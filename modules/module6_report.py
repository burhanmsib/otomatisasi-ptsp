from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
from docx.text.paragraph import Paragraph
from docx.enum.table import WD_ALIGN_VERTICAL
from datetime import datetime
from io import BytesIO
import re

# =========================
# CONSTANTS
# =========================
ID_MONTHS = {
    1: "Januari", 2: "Februari", 3: "Maret", 4: "April", 5: "Mei", 6: "Juni",
    7: "Juli", 8: "Agustus", 9: "September", 10: "Oktober",
    11: "November", 12: "Desember"
}

# =========================
# DATE PARSER
# =========================
def parse_date_flexible(date_str: str):
    if not date_str:
        return None

    s = str(date_str).strip()

    month_map = {
        "Januari": "January", "Februari": "February", "Maret": "March", "April": "April",
        "Mei": "May", "Juni": "June", "Juli": "July", "Agustus": "August",
        "September": "September", "Oktober": "October", "November": "November", "Desember": "December"
    }
    for indo, eng in month_map.items():
        s = s.replace(indo, eng)

    fmts = [
        "%d %B %Y", "%d.%m.%Y", "%d-%m-%Y",
        "%Y-%m-%d", "%d/%m/%Y", "%d %b %Y", "%m/%d/%Y"
    ]

    for f in fmts:
        try:
            return datetime.strptime(s, f)
        except:
            pass

    m = re.search(r"(\d{1,2})[./-](\d{1,2})[./-](\d{2,4})", s)
    if m:
        d, mn, y = m.groups()
        if len(y) == 2:
            y = "20" + y
        try:
            return datetime.strptime(f"{d}-{mn}-{y}", "%d-%m-%Y")
        except:
            pass

    m2 = re.search(r"(\d{4})[./-](\d{1,2})[./-](\d{1,2})", s)
    if m2:
        y, mn, d = m2.groups()
        try:
            return datetime.strptime(f"{d}-{mn}-{y}", "%d-%m-%Y")
        except:
            pass

    return None


def format_date_id(dt: datetime):
    if not dt:
        return ""
    return f"{dt.day:02d} {ID_MONTHS.get(dt.month, '')} {dt.year}"


def format_date_en(dt: datetime):
    if not dt:
        return ""

    return (
        f"{dt.strftime('%B')} "
        f"{dt.day}, "
        f"{dt.year}"
    )


# =========================
# STYLE UTILITIES
# =========================
def style_paragraph(
    p,
    size=12,
    bold=False,
    italic=False,
    align="left",
    space_before=0,
    space_after=0,
    line_spacing=1.0,
    left_indent_cm=0,
    first_line_indent_cm=0
):
    if not p.runs:
        p.add_run("")

    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        try:
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        except:
            pass

    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    pf.left_indent = Cm(left_indent_cm)
    pf.first_line_indent = Cm(first_line_indent_cm)


def set_table_border(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)

    borders = OxmlElement("w:tblBorders")
    for side in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        elem = OxmlElement(f"w:{side}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:color"), "000000")
        borders.append(elem)

    tblPr.append(borders)


# =========================
# XML HELPERS
# =========================
def insert_paragraph_after(paragraph, text=None, style=None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)

    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style

    return new_para


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def delete_paragraph(paragraph):
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def remove_template_markers(doc):
    markers = {
        "$LAPORAN_SECTION_START",
        "$LAPORAN_SECTION_END",
        "$Laporan_Section_Start",
        "$Laporan_Section_End",
    }

    paragraphs_to_delete = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if text in markers:
            paragraphs_to_delete.append(p)

    for p in reversed(paragraphs_to_delete):
        delete_paragraph(p)


# =========================
# SECTION BUILDERS
# =========================
def build_title(doc, row):

    dt = parse_date_flexible(row.get("Tanggal Koordinat", ""))

    t_str = (
        format_date_en(dt)
        if dt
        else row.get("Tanggal Koordinat", "")
    )
    # 🔥 FIX: pakai format asli (DMS)
    coord_text = row.get("Koordinat", "")
    
    ka = str(row.get("Koordinat Awal", "") or "").strip()
    kb = str(row.get("Koordinat Akhir", "") or "").strip()
    
    if ka and kb:
        if ka == kb:
            coord_text = ka
        else:
            coord_text = f"{ka} to {kb}"
    else:
        coord_text = "-"
    # ka = row.get("Koordinat Awal", "")
    # kb = row.get("Koordinat Akhir", "")

    # # 🔥 FIX TITIK VS RUTE
    # if ka and kb:
    #     if str(ka) == str(kb):
    #         coord_text = str(ka)
    #     else:
    #         coord_text = f"{ka} to {kb}"
    # else:
    #     coord_text = "-"

    p = doc.add_paragraph()
    p.add_run("Meteorological Reports\n").bold = True
    p.add_run("Coordinate from ").bold = True
    p.add_run(f"{coord_text}\n")
    p.add_run(f"for {t_str}")

    style_paragraph(p, bold=True, align="center")

    doc.add_paragraph("")
    
def build_interval_table(doc, intervals, tz="WIB"):
    headers = [
        "DATE", f"LOCAL TIME ({tz})", "WEATHER",
        "WIND (Knot)", "CURRENT (cm/s)",
        "WAVE (meter)", "BEAUFORT SCALE"
    ]

    table = doc.add_table(rows=1, cols=7)
    set_table_border(table)

    # HEADER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        style_paragraph(cell.paragraphs[0], bold=True, align="center")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    # =========================
    # ADD DATA ROWS
    # =========================
    for j in range(4):
        data = intervals[j] if j < len(intervals) else {}
        row = table.add_row().cells

        values = [
            "",  # DATE dikosongkan dulu
            data.get("LOCAL TIME", ""),
            data.get("WEATHER", ""),
            data.get("WIND", ""),
            data.get("CURRENT", ""),
            data.get("WAVE", ""),
            data.get("BEAUFORT", ""),
        ]

        for i, v in enumerate(values):
            cell = row[i]
            cell.text = str(v)
            style_paragraph(cell.paragraphs[0], align="center")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # =========================
            # MERGE DATE + CENTER (SAFE)
            # =========================
            if intervals and len(table.rows) > 1:
            
                date_text = intervals[0].get("DATE", "")
            
                # row pertama setelah header = index 1
                start_row = 1
                end_row = len(table.rows) - 1  # terakhir
            
                start_cell = table.cell(start_row, 0)
                end_cell = table.cell(end_row, 0)
            
                merged_cell = start_cell.merge(end_cell)
                merged_cell.text = date_text
            
                # Center horizontal
                style_paragraph(merged_cell.paragraphs[0], align="center")
            
                # Bold
                if merged_cell.paragraphs[0].runs:
                    merged_cell.paragraphs[0].runs[0].bold = True
            
                # Center vertical
                merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def build_notes_primary(doc):

    p0 = doc.add_paragraph()
    run0 = p0.add_run("*) satellite images enclosed")
    run0.italic = True

    style_paragraph(
        p0,
        size=11
    )

    p = doc.add_paragraph()

    r1 = p.add_run("Note:\n")
    r1.bold = True
    r1.italic = True

    r2 = p.add_run("The direction of current is ")

    r3 = p.add_run("toward")
    r3.bold = True

    p.add_run(".\n")

    r4 = p.add_run("The direction of wind is ")

    r5 = p.add_run("from")
    r5.bold = True

    p.add_run(".")

    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)

    p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def build_wave_category_table(doc):
    p = doc.add_paragraph()

    run = p.add_run("Note :")
    run.bold = True
    
    style_paragraph(
        p,
        size=11
    )
    data = [
        ("Smooth", "0.10 – 0.50 m"),
        ("Slight", "0.50 – 1.25 m"),
        ("Moderate", "1.25 – 2.50 m"),
        ("Rough", "2.50 – 4.00 m"),
        ("Very Rough", "4.00 – 6.00 m"),
        ("High", "6.00 – 9.00 m"),
        ("Very High", "9.00 – 14.00 m"),
    ]

    t = doc.add_table(rows=0, cols=2)
    t.autofit = False
    set_table_border(t)

    for label, val in data:

        row = t.add_row()
    
        row.cells[0].width = Cm(2.6)
        row.cells[1].width = Cm(2.6)
    
        row.cells[0].text = label
        row.cells[1].text = val
    
        style_paragraph(
            row.cells[0].paragraphs[0],
            size=10,
            align="left"
        )
    
        style_paragraph(
            row.cells[1].paragraphs[0],
            size=10,
            align="center"
        )

    doc.add_paragraph("")


def build_satellite_image_table(doc, tanggal_str):
    dt = parse_date_flexible(tanggal_str)
    tanggal_fmt = format_date_en(dt) if dt else tanggal_str

    table = doc.add_table(rows=2, cols=2)
    set_table_border(table)

    doc.add_paragraph("")
    doc.add_paragraph("")

    hdr = table.rows[0].cells[0]
    hdr.merge(table.rows[0].cells[1])

    p = hdr.paragraphs[0]
    p.add_run(f"Weather Satellite Image on {tanggal_fmt} at ______")
    style_paragraph(p, bold=True, align="center")

    table.rows[1].cells[0].paragraphs[0].add_run("[Insert Satellite Image Here]")
    table.rows[1].cells[1].paragraphs[0].add_run("[Insert Legend Here]")

    style_paragraph(table.rows[1].cells[0].paragraphs[0], italic=True, align="center")
    style_paragraph(table.rows[1].cells[1].paragraphs[0], italic=True, align="center")

    doc.add_paragraph("")


# =========================
# FIRST PAGE PLACEHOLDER REPLACER (FINAL FULL FIX - NO LOGIC REMOVED)
# =========================
def replace_first_page_placeholders(doc, module1_rows, module5_rows):

    # =========================
    # 🔥 BASE (TETAP PAKAI LOGIC ASLI)
    # =========================
    first = module1_rows[0]

    # =========================
    # 🔥 TAMBAHAN: FALLBACK (TANPA MERUBAH FLOW)
    # =========================
    def get_fallback(key):
        val = str(first.get(key, "") or "").strip()

        if not val or val == "-":
            for r in module1_rows:
                v = str(r.get(key, "") or "").strip()
                if v and v != "-":
                    return v
            return ""
        return val

    nama_perusahaan = get_fallback("Nama Perusahaan")
    alamat_perusahaan = get_fallback("Alamat Perusahaan")
    ref_no = get_fallback("Nomor Surat")

    if not ref_no:
        ref_no = "______"

    # =========================
    # LOGIC ASLI (TIDAK DIUBAH)
    # =========================
    valid_report_count = sum(
        1
        for idx in range(len(module1_rows))
        if idx < len(module5_rows)
        and module5_rows[idx]
        and "intervals" in module5_rows[idx]
    )

    replacements = {
        "$nama_perusahaan": nama_perusahaan,
        "$alamat_perusahaan": alamat_perusahaan,
        "$no_surat": ref_no,
        "$tanggal_hari_ini": format_date_en(datetime.now()),
        "$jumlah_laporan_section": str(valid_report_count),
    }

    paragraphs_to_delete = []

    for p in doc.paragraphs:
        text = p.text.strip()

        if text.startswith("Responding to your letter") and "$LIST_KOORDINAT" not in text:
            paragraphs_to_delete.append(p)
            continue

        if "here with we enclose the meteorological analysis" in text.lower():
            paragraphs_to_delete.append(p)
            continue

        # =========================
        # 🔥 BAGIAN KOORDINAT (LOGIC ASLI + UPGRADE)
        # =========================
        if "$LIST_KOORDINAT" in text:
            clear_paragraph(p)

            intro_text = (
                f"Responding to your letter with Ref. {ref_no} "
                f"on the subject of marine meteorological analysis with coordinate :"
            )
            p.add_run(intro_text)

            style_paragraph(
                p,
                size=12,
                align="justify",
                space_before=0,
                space_after=2,
                line_spacing=1.0
            )

            current_p = p

            # =========================
            # 🔥 GROUP BY KOORDINAT (TAMBAHAN)
            # =========================
            coord_map = {}

            for row in module1_rows:

                ka = str(row.get("Koordinat Awal", "") or "").strip()
                kb = str(row.get("Koordinat Akhir", "") or "").strip()
                dt = parse_date_flexible(row.get("Tanggal Koordinat", ""))
            
                if not ka or not kb:
                    continue
            
                dt_str = format_date_en(dt) if dt else ""
            
                bullet_text = (
                    f"• From {ka} to {kb} for {dt_str}"
                )
            
                new_p = insert_paragraph_after(current_p)
                clear_paragraph(new_p)
                new_p.add_run(bullet_text)
            
                style_paragraph(
                    new_p,
                    size=11,
                    align="justify",
                    space_before=0,
                    space_after=0,
                    line_spacing=1.0,
                    left_indent_cm=0.5,
                    first_line_indent_cm=-0.3
                )
            
                current_p = new_p

            # =========================
            # 🔥 BUILD BULLET (FORMAT ASLI)
            # =========================
            for coord, dates in coord_map.items():

                if not dates:
                    continue

                dates = sorted(dates)

                if len(dates) == 1:
                    dt_str = format_date_en(dates[0])
                else:
                    start = dates[0]
                    end = dates[-1]

                    if start.month == end.month and start.year == end.year:
                        dt_str = f"{start.strftime('%B %d')}–{end.strftime('%d, %Y')}"
                    else:
                        dt_str = f"{format_date_en(start)} – {format_date_en(end)}"

                bullet_text = f"• {coord} for {dt_str}"

                new_p = insert_paragraph_after(current_p)
                clear_paragraph(new_p)
                new_p.add_run(bullet_text)

                style_paragraph(
                    new_p,
                    size=11,
                    align="justify",
                    space_before=0,
                    space_after=0,
                    line_spacing=1.0,
                    left_indent_cm=0.5,
                    first_line_indent_cm=-0.3
                )

                current_p = new_p

            # =========================
            # PENUTUP (TETAP)
            # =========================
            end_p = insert_paragraph_after(current_p)
            clear_paragraph(end_p)
            end_p.add_run("here with we enclose the meteorological analysis in attachments sheets.")

            style_paragraph(
                end_p,
                size=12,
                align="justify",
                space_before=2,
                space_after=0,
                line_spacing=1.0
            )

            continue

        # =========================
        # REPLACEMENT TEXT (TETAP)
        # =========================
        for k, v in replacements.items():
            if k in p.text:
                p.text = p.text.replace(k, str(v))
                style_paragraph(p)

    for p in reversed(paragraphs_to_delete):
        delete_paragraph(p)

# =========================
# MAIN ENTRY
# =========================
def generate_final_docx_streamlit(module1_rows, module5_rows, template_path):
    doc = Document(template_path)

    replace_first_page_placeholders(doc, module1_rows, module5_rows)

    for idx, row in enumerate(module1_rows):
        if not module5_rows:
            continue
        if idx >= len(module5_rows):
            continue

        module5_item = module5_rows[idx]
        if module5_item is None:
            continue
        if "intervals" not in module5_item:
            continue

        build_title(doc, row)

        intervals = module5_item["intervals"]
        tz = module5_item.get("tz", "WIB")

        build_interval_table(doc, intervals, tz)
        build_notes_primary(doc)
        build_wave_category_table(doc)
        build_satellite_image_table(doc, row.get("Tanggal Koordinat", ""))

        if idx < len(module1_rows) - 1:
            doc.add_page_break()

    remove_template_markers(doc)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
